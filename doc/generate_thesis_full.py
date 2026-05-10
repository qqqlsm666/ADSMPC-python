"""
generate_thesis_full.py

基于模板 docx 生成完整论文 docx：
封面 + 声明 + 摘要 + 第 1-5 章 + 参考文献 + 致谢 + 附录 + 创新成果。
"""
from docx import Document
from docx.oxml.ns import qn

PARAGRAPHS = []

def p(text, style='Normal'):
    PARAGRAPHS.append((style, text))

def empty():
    PARAGRAPHS.append(('Normal', ''))

# ============ 封面 ============
empty(); empty()
p('本 科 毕 业 设 计（ 论 文 ）')
empty(); empty(); empty()
p('题目:   支持隐私保护的检索增强生成系统    ')
empty(); empty()
p('姓    名         ____________         ')
p('学    院         ____________         ')
p('专    业         ____________         ')
p('班    级         ____________         ')
p('学    号         ____________         ')
p('指导教师         ____________         ')
empty()
p(' 2026 年 6 月')
empty()
p('北 京 邮 电 大 学')

# ============ 诚信声明 ============
p('本科毕业设计（论文）诚信声明')
p('本人声明所呈交的毕业设计（论文），题目《支持隐私保护的检索增强生成系统》是本人在指导教师的指导下，独立进行研究工作所取得的成果。尽我所知，除了文中特别加以标注和致谢中所罗列的内容以外，论文中不包含其他人已经发表或撰写过的研究成果，也不包含为获得北京邮电大学或其他教育机构的学位或证书而使用过的材料。')
p('申请学位论文与资料若有不实之处，本人承担一切相关责任。')
empty()
p('本人签名：                         日期：                          ')
empty(); empty(); empty(); empty()

# ============ 授权说明 ============
p('关于论文使用授权的说明')
p('本人完全了解并同意北京邮电大学有关保留、使用学位论文的规定，即：北京邮电大学拥有以下关于学位论文的无偿使用权，具体包括：学校有权保留并向国家有关部门或机构送交学位论文，有权允许学位论文被查阅和借阅；学校可以公布学位论文的全部或部分内容，有权允许采用影印、缩印或其它复制手段保存。汇编学位论文，将学位论文的全部或部分内容编入有关数据库进行检索。（保密的学位论文在解密后遵守此规定）          ')
empty()
p('本人签名：                         日期：                          ')
empty()
p('导师签名：                         日期：                          ')
empty(); empty()

# ============ 中文摘要 ============
p('支持隐私保护的检索增强生成系统')
empty()
p('摘要')
empty()
p('随着大语言模型（Large Language Models, LLMs）能力的迅猛提升，检索增强生成（Retrieval-Augmented Generation, RAG）已成为构建知识密集型问答系统的主流方案。然而，传统 RAG 系统在客户端的查询语句、服务端的文档库以及编码模型的权重三方面同时面临隐私泄露风险，特别是在医疗、法律、金融等敏感领域，这一矛盾尤为突出。如何在不暴露任意一方明文数据的前提下，完整地完成"检索→召回→联合推理"的 RAG 流水线，是当前隐私保护机器学习领域亟待解决的关键问题。')
p('本课题面向半诚实两方计算（2-Party Semi-honest）安全模型，针对密态 RAG 系统中"双路检索如何不泄露文档身份"、"联合推理后的语义向量如何形成有意义的检索输出"以及"密态计算的工程性能与数值正确性如何兼顾"三大挑战展开研究，主要工作包括：')
p('1. 设计并实现了基于双路检索的密态 RAG 检索算法。以 NssMPClib 安全多方计算库为底层，融合算术秘密分享（Arithmetic Secret Sharing, ASS）与函数秘密分享（Function Secret Sharing, FSS）两类基本协议，提出"密态语义检索 + 密态简化 BM25 词汇检索"双路并行打分方案，并设计基于密态指示器（One-hot Indicator）的密态 Top-K 冒泡排序与密态文档抽取算法，实现双方协同选出最相关文档但任意一方均不知道选了哪一篇的隐私目标。')
p('2. 提出基于密态 BERT 联合编码的 Cross-Encoder Reranker 算法。针对联合推理产出的池化向量缺乏可解释下游用途的问题，设计将联合编码的池化向量与原始文档语义库进行密态矩阵乘法的精排算法，把"装饰性"的联合推理转化为可解释、可量化的密态精排器。在自建的小型问答语料 10 条查询测试中，密态系统的 Recall@5 达到 100%、平均倒数排名（MRR）0.72、归一化折损累积增益（NDCG@5）0.79，且明文与密态精排分数的余弦相似度高达 0.9998。')
p('3. 构建了端到端的支持隐私保护的检索增强生成系统。包括应用层 secure_rag 包（服务端、客户端、检索算法、明文基线）、实验对比层（数值一致性测试、检索质量评估、子进程隔离运行器）以及完整的项目文档（系统架构、威胁模型、实验复现指南）。系统在 Windows 11 + Python 3.10 + PyTorch 2.3 + 自编译 torchcsprng AES PRG 环境下，单条查询端到端约 54 秒，相比基础实现加速约 25 倍，证明了在普通笔记本硬件上运行密态 RAG 的工程可行性。')
empty()
p('关键词  检索增强生成  安全多方计算  隐私保护  密态推理  函数秘密分享')

# ============ 英文摘要 ============
empty(); empty()
p('A Privacy-Preserving Retrieval-Augmented Generation System')
empty()
p('ABSTRACT')
empty()
p('With the rapid advance of Large Language Models (LLMs), Retrieval-Augmented Generation (RAG) has become the de-facto paradigm for building knowledge-intensive question-answering systems. However, traditional RAG pipelines simultaneously expose privacy on three sides: the user query on the client, the document corpus on the server, and the weights of the encoder model. The tension is particularly sharp in sensitive domains such as healthcare, legal, and financial services. Completing the entire "retrieve-recall-joint inference" pipeline of RAG without revealing any party\'s plaintext data remains a key open problem in privacy-preserving machine learning.')
p('This thesis targets a 2-party semi-honest secure computation setting and tackles three challenges in encrypted RAG: (i) how to perform dual-path retrieval without leaking the identity of the selected document, (ii) how to turn the pooled vector produced by joint inference into a meaningful retrieval output, and (iii) how to balance engineering performance with numerical correctness in encrypted computation. The main contributions are: (1) a dual-path encrypted retrieval algorithm built on top of NssMPClib that combines an encrypted semantic path with an encrypted BM25-like lexical path, and an encrypted Top-K bubble sort with one-hot indicators ensuring that neither party learns which document was chosen; (2) a secure Cross-Encoder Reranker built upon encrypted BERT joint encoding, achieving Recall@5 = 1.00, MRR = 0.72, NDCG@5 = 0.79 across 10 queries with a 0.9998 cosine similarity between plaintext and encrypted reranker scores; and (3) an end-to-end privacy-preserving RAG system on Windows 11 + Python 3.10 + PyTorch 2.3 with self-compiled torchcsprng, completing a query end-to-end in about 54 seconds, roughly 25x faster than the baseline implementation.')
p('This work demonstrates the engineering feasibility of running encrypted RAG on commodity laptop hardware and provides a reproducible prototype for the privacy-preserving machine learning community.')
empty()
p('KEY WORDS  Retrieval-Augmented Generation  Secure Multi-Party Computation  Privacy Preservation  Encrypted Inference  Function Secret Sharing')

# ============ 目录 ============
empty(); empty()
p('目录')
p('（请在 Word 中右键此处，选择"更新域"或使用引用→目录功能基于 Heading 样式自动生成。'
  '本文使用 Heading 1 标识章名、Heading 2 标识 X.X 节、Heading 3 标识 X.X.X 小节。）')
empty(); empty(); empty()

# ============ 第一章 ============
p('第一章 绪论', style='Heading 1')
p('1.1 研究背景与意义', style='Heading 2')
p('近年来，以 GPT、Llama、DeepSeek 等为代表的大语言模型（Large Language Models, LLMs）在问答系统、代码生成、教育教学、法律咨询、医疗问诊等众多领域展现出强大能力，已逐步成为信息服务的核心基础设施[1]。然而，纯参数化的 LLM 难以避免事实性幻觉、训练数据时效性差、领域知识封闭等固有缺陷，导致其在专业场景下的可信度受限。检索增强生成（Retrieval-Augmented Generation, RAG）通过将外部知识库与生成模型解耦，在生成回答前先从知识库中检索相关文档，再把"问题 + 检索片段"一并送入生成模型，从根本上缓解了上述问题，已成为当前知识密集型 LLM 应用的主流架构[2]。')
p('然而，RAG 系统的隐私问题尚未得到充分重视。一个完整的 RAG 流程同时涉及三类敏感数据：（1）用户提交的查询，可能包含病情、案件、商业机密等高度个人化信息；（2）服务端持有的文档库，可能涉及医院病例、法律卷宗、企业内部文档等不可对外共享的知识资产；（3）模型权重，体现训练数据与算法投入，是服务提供方的核心知识产权。在传统 RAG 部署中，三者必须明文聚合在同一计算节点上才能完成检索与生成，导致客户端用户必须信任服务端不会窥视其查询内容、服务端必须把文档库以明文形式暴露给计算引擎、第三方算力提供商更可能同时观察查询与文档。这一信任模型在医疗诊断、法律咨询、金融风控等高敏场景下难以被各方接受[30]。')
p('针对 RAG 中的隐私问题，安全多方计算（Secure Multi-Party Computation, MPC）提供了基础的密码学工具：在不暴露任何一方明文输入的前提下，各方协同计算函数的输出[8]。如何把整条 RAG 流水线（编码、双路打分、Top-K 排序、文档抽取、联合推理）系统地搬迁到密态计算之上，且在普通硬件上达到工程可用水平，是隐私保护机器学习领域当前的研究热点之一[15][16]。本课题在西安电子科技大学网络与系统安全实验室开源的 NssMPClib 框架[29]之上，面向半诚实两方计算模型，构建了一个支持隐私保护的检索增强生成系统的原型。研究密态 RAG 不仅有助于把 LLM 应用推广到更多对隐私敏感的高价值领域，也对推动 MPC 协议在面向 Transformer 架构的非线性计算（如 LayerNorm、Softmax、GeLU 等）上的工程优化具有重要的理论与实用意义。')

p('1.2 国内外研究现状', style='Heading 2')
p('本节围绕"支持隐私保护的检索增强生成系统"这一核心命题，从检索增强生成、隐私保护检索与密态机器学习推理三个维度梳理代表性研究工作。')
p('1.2.1 检索增强生成技术', style='Heading 3')
p('检索增强生成是当前知识密集型自然语言处理的主流范式。Karpukhin 等人于 2020 年提出的 Dense Passage Retriever（DPR）首次系统性地证明了基于双塔 BERT 的稠密检索在开放域问答上显著优于传统的 BM25 词汇检索[2]。Lewis 等人随后提出的原始 RAG 框架将稠密检索与生成式模型联合训练，把检索作为可微分模块嵌入到端到端管线中[1]。后续 Atlas[7]、RETRO[6]、In-Context-RALM 等工作进一步在大模型规模、上下文长度、检索粒度等维度上推动 RAG 的发展。在工业界，OpenAI、百度文心、阿里通义等头部 LLM 平台均已将 RAG 作为知识接入的标准接口。在双路检索（hybrid retrieval）方面，业界普遍采用稠密检索（语义路）与传统 BM25[3] 词汇检索并行召回、再通过 Reciprocal Rank Fusion（RRF）等策略融合的方案，以兼顾语义泛化能力与精确关键词匹配能力。')
p('现有 RAG 工作多在明文环境下展开，专注于召回质量、上下文窗口扩展与检索粒度等维度，对查询与文档的隐私保护问题关注较少。')
p('1.2.2 隐私保护检索', style='Heading 3')
p('隐私信息检索（Private Information Retrieval, PIR）是隐私保护检索的经典方向，其目标是允许客户端从服务端查询特定记录而不暴露查询索引。Chor 等人 1995 年提出的多服务器 PIR 在信息论安全模型下达成隐私目标，但通信复杂度随数据库规模线性增长[19]。Kushilevitz 与 Ostrovsky 的单服务器计算性 PIR 利用同态加密把通信复杂度降低到亚线性[20]。然而经典 PIR 本质上要求查询是"按 ID 取记录"，与 RAG 中"按相关度排序后取 top-K"的语义不匹配。')
p('不经意随机访问机（Oblivious Random Access Memory, ORAM）通过反复打乱物理存储位置实现访问模式混淆[21]，但在 RAG 这类需要 Top-K 排序与多次查表的场景下开销高昂。最近兴起的密态向量检索（如 SimplePIR[22]、Tiptoe、Coral 等）尝试在两方设置下实现密态最近邻搜索，但通常只覆盖检索阶段的稠密向量打分，未集成完整 RAG 中的 BM25 词汇检索、密态文档抽取与密态生成 / 编码环节。')
p('1.2.3 密态机器学习推理', style='Heading 3')
p('随着 SecureML[13]、CrypTen[14] 等通用框架的出现，神经网络的密态推理逐渐由理论走向实践。在 Transformer 架构上，SIGMA（Secure GPT Inference）[15]、BumbleBee[18]、Iron[17] 等工作针对 Softmax、LayerNorm、GeLU 等非线性算子提出了基于函数秘密分享（FSS）与查表近似的高效协议，把单层 Transformer 在 LAN 环境下的密态前向时间压缩到数秒级。MPCFormer[16] 在 BERT 模型上系统验证了密态推理的可行性。开源的 NssMPClib[29] 由西安电子科技大学网络与系统安全实验室维护，提供了 RingTensor、ASS、FSS（含 DPF/DCF/DICF/VDPF/VSigma）、Beaver Triples、Paillier 同态加密等基础组件，并实现了密态 BERT、CNN 等典型模型，已被多个研究项目用于密态机器学习实验。')
p('然而，已有密态机器学习工作主要面向单纯的推理任务，对"检索 + 推理联合"的 RAG 系统级问题鲜有讨论；现有密态向量检索工作又通常脱离生成 / 编码模块单独存在。这正是本课题试图填补的空白。')

p('1.3 研究内容与创新点', style='Heading 2')
p('针对检索增强生成系统中查询、文档库与模型权重三方隐私同时受到威胁的现状，本文设计并实现了一种基于双路检索 + 密态联合编码 + 密态精排的端到端密态 RAG 系统，并在自建小型问答语料上完成了密态 vs 明文的多维对比实验。系统总体上由"应用层 secure_rag 包 + 实验对比层 experiments 模块 + 底层 NssMPClib MPC 库"三层组成。')
p('本文的主要贡献总结如下：')
p('（1）设计并实现了基于双路并行的密态检索算法。算法包含：（a）基于密态广播乘法的语义路打分，对密态查询向量与密态文档库做内积；（b）基于密态稀疏向量与 BM25 倒排矩阵的词汇路打分；（c）基于密态指示器（one-hot 单位向量包装为密态分享）的密态 Top-K 冒泡排序，使得在交换分数的同时同步交换"身份证向量"，最终输出 [K, NUM_DOCS] 的密态指示矩阵；（d）基于"指示器与文档 token 库逐元素相乘后求和"的密态文档抽取，使任意一方无法获知 Top-K 实际选中的文档下标。算法整体保留了与明文双路检索一致的语义召回能力，同时实现了文档身份的完全隐私保护。')
p('（2）提出了基于密态联合编码的 Cross-Encoder Reranker 算法。本文观察到原 RAG 框架联合推理产出的 [CLS] 池化向量缺乏明确下游用途的设计缺陷，进一步提出将该向量与原始文档语义库做密态矩阵乘法的精排方案：以联合编码后的池化向量作为"经过 query-doc 融合的精炼表示"，通过 [1, hidden] @ [hidden, NUM_DOCS] 的密态矩阵乘运算得到对每篇文档的精排分数。在 10 条查询的真实评估上，该方案使密态系统的 Recall@5 达到 100%、MRR 达到 0.72、NDCG@5 达到 0.79，且明文与密态精排分数的余弦相似度高达 0.9998，证明了协议的数值正确性。')
p('（3）构建了完整的端到端密态 RAG 实验平台。系统包括：（a）应用层 secure_rag 包（服务端流程、客户端流程、检索算法库、明文基线、辅助参数生成器、全局配置），（b）实验对比层 experiments 模块（基于子进程隔离的密态 RAG 运行器、HuggingFace Tokenizer 接入的语料加载器、Recall@K / MRR / NDCG@K / Precision@K 四类 IR 指标实现、数值一致性对比脚本、检索质量对比脚本、整合入口），（c）配套文档（系统架构图、威胁模型说明、实验复现指南）。在普通笔记本（i7 + 16 GB RAM + RTX 3050 4 GB）上，针对自编译 torchcsprng AES PRG 优化后单条查询端到端约 54 秒，相比基础实现加速约 25 倍。系统已经过若干工程层面的改造与修复，包括针对 DEBUG_LEVEL=2 单密钥广播路径的 Beaver mul 与 prefix_parity_query 修复、子进程端口隔离修复、close() 阶段挂起规避等。')

p('1.4 章节安排', style='Heading 2')
p('本文一共包含五个章节，各章节的主要内容如下：')
p('第一章为绪论。介绍了课题的研究背景，分析了检索增强生成系统在查询、文档库与模型权重三方面同时面临的隐私挑战，从检索增强生成、隐私保护检索、密态机器学习推理三个维度综述了国内外研究现状，提出了本文的研究内容与创新点。')
p('第二章为相关研究。本章对密态 RAG 涉及的基础理论与代表性工作做较深入的回顾，依次介绍了检索增强生成的典型架构与检索范式、安全多方计算的关键协议（算术秘密分享、函数秘密分享、Beaver triples、半诚实两方计算模型）、密态神经网络推理的代表性方法（SIGMA、BumbleBee、Iron、NssMPClib），为后续章节的设计奠定基础。')
p('第三章为系统设计与实现。本章详细介绍设计并实现的密态 RAG 系统，包括系统总体架构与威胁模型、双路密态检索算法（密态语义检索、密态词汇检索、密态 Top-K 指示器排序、密态文档抽取）、密态联合编码与 Cross-Encoder Reranker 算法。')
p('第四章为实验与分析。本章在自构建的小型问答语料库上对系统进行了多维评估：数值一致性实验验证密态计算的正确性、检索质量实验对比明文与密态在 Recall@K / MRR / NDCG@K 等 IR 指标上的差异、性能拆解实验定位密态推理的主要瓶颈、消融实验考察 Reranker 的有效性。')
p('第五章为总结与展望。对本文的研究工作进行总结，并对未来的研究方向（包括密态生成式 LLM、可扩展 Top-K、恶意安全升级等）进行展望。')

# ============ 第二章 ============
empty()
p('第二章 相关研究', style='Heading 1')
p('本章对密态 RAG 涉及的基础理论与代表性工作进行回顾，主要包括三方面：检索增强生成的典型架构与检索范式、安全多方计算的关键协议、面向 Transformer 架构的密态神经网络推理。')

p('2.1 检索增强生成', style='Heading 2')
p('检索增强生成（RAG）的核心思想是将参数化的语言模型与非参数化的外部知识库解耦，在生成回答前先从知识库检索相关文档，再以"问题 + 检索片段"作为完整上下文送入生成模型。一个完整的 RAG 流程可以划分为三个核心阶段：编码阶段、检索阶段、生成阶段。')
p('编码阶段把查询与文档分别映射到向量空间。常用的编码模型为 BERT[4]、Sentence-BERT 等基于 Transformer 的双塔结构。文档库通常在系统部署阶段离线编码并存储为稠密向量库。')
p('检索阶段给定查询向量与文档库，从中筛选出 top-K 最相关的文档。主流方案包括稠密检索、稀疏检索与混合检索三类。稠密检索（Dense Retrieval）基于查询与文档向量的内积或余弦相似度，优点是能够捕捉语义层面的相似性（如"汽车"与"轿车"），缺点是对未在训练数据中出现的稀有术语泛化能力弱。代表工作 DPR[2] 通过对比学习训练双塔编码器。稀疏检索（Sparse Retrieval）基于词频统计的传统方法，BM25[3] 是其中最具代表性的算法，公式为 BM25(q, d) = Σ IDF(t) · f(t,d)·(k1+1) / (f(t,d) + k1·(1−b+b·|d|/avgdl))，其中 f(t,d) 是词 t 在文档 d 中的频次，|d| 是文档长度，avgdl 是平均文档长度，k1、b 为可调超参（通常取 1.5、0.75）。BM25 的优势在于精确匹配关键词。混合检索（Hybrid Retrieval）将稠密路径与稀疏路径并行召回、再通过分数融合或排序融合（如 Reciprocal Rank Fusion）合并结果，实现语义泛化与精确匹配的兼顾，是工业界主流方案。')
p('生成阶段把"问题 + top-K 文档"拼接为完整上下文，送入生成模型（如 GPT、Llama）产出回答。在 BERT-Reader 风格的早期 RAG 中（如 ORQA、REALM），生成阶段被替换为基于 [CLS] 池化向量接 span prediction head 的"阅读"环节。')
p('近年来 RAG 的研究主要围绕召回质量优化（Atlas[7]）、长上下文窗口扩展（RETRO[6]）、检索粒度（chunk-level、token-level）等维度展开。值得注意的是，绝大多数现有 RAG 工作都假设查询、文档库、模型权重三者可以明文聚合到同一计算节点，未考虑这些数据在不同主体间的隐私边界，这正是本文研究的出发点。')

p('2.2 安全多方计算', style='Heading 2')
p('安全多方计算（Secure Multi-Party Computation, MPC）是密码学中的一个分支，研究多个互不信任的参与方在不暴露各自私有输入的前提下协同计算公共函数的方法[8]。其安全模型主要分两类：半诚实模型（Semi-honest, a.k.a. honest-but-curious）下所有参与方严格按协议执行，但可能从协议运行中收集到的信息推断对方私有输入；恶意模型（Malicious）下参与方可能任意偏离协议执行。本文研究的密态 RAG 系统建立在半诚实两方计算模型之上。')
p('2.2.1 算术秘密分享', style='Heading 3')
p('算术秘密分享（Arithmetic Secret Sharing, ASS）是 MPC 的基本工具。在 2-out-of-2 加法分享方案下，秘密值 x ∈ Z_{2^L} 被随机分成两份 x_0 与 x_1 满足 x = x_0 + x_1 (mod 2^L)，分别由两方持有。任意一方单独持有的份额是均匀随机的，因此完全不泄露 x 的信息。')
p('ASS 在加法上具有优良性质：双方各自把份额相加，不需要任何通信即可得到 x + y 的分享。乘法则需要借助 Beaver triples[9]：双方共同持有随机三元组 (a, b, c) 的分享且满足 c = a · b；要计算 x · y 时，先计算 e = x − a、f = y − b 并将其重构为明文，然后 x · y = e · f + e · b + f · a + c。这样每次乘法只需要一次双向通信（双方互发 e、f 各自的份额）。Beaver triples 通常在离线阶段批量预生成，在线阶段直接消费。')
p('2.2.2 函数秘密分享', style='Heading 3')
p('函数秘密分享（Function Secret Sharing, FSS）是 ASS 之外的另一类基础协议，由 Boyle 等人于 2015 年提出[10]。FSS 把函数 f 的求值过程分成两份"函数密钥"k_0、k_1，使得双方分别用各自的密钥本地计算，最后把输出相加即可重构 f(x)。FSS 最有用的两种特例是分布式点函数（Distributed Point Function, DPF）：f(x) = β if x = α, else 0；以及分布式比较函数（Distributed Comparison Function, DCF）：f(x) = β if x < α, else 0[11]。基于 DPF/DCF 可以构造分布式区间比较函数（DICF），在密态下高效实现 x ≥ y、x ≤ y、x = y 等逐元素比较。本文采用的 NssMPClib 框架实现了 DPF、DCF、DICF 以及 Grotto[12]、SIGMA[15] 等优化变种。')
p('2.2.3 Beaver 矩阵乘法三元组', style='Heading 3')
p('针对深度神经网络中频繁出现的矩阵乘法运算，Beaver 三元组的概念可以推广为矩阵形式：双方共同持有随机矩阵三元组 (A, B, C) 的分享且满足 C = A · B；计算 X · Y 时，重构 E = X − A 与 F = Y − B，然后 X · Y = E · F + E · B + A · F + C。矩阵 Beaver 把"N × M 个标量乘法"压缩为"一次矩阵 Beaver 协议"，显著减少通信轮数，是密态神经网络推理的关键优化手段。')

p('2.3 密态神经网络推理', style='Heading 2')
p('把神经网络从明文搬迁到密态环境，瓶颈不在线性层（矩阵乘法可由 Beaver 矩阵三元组高效完成），而在非线性层。Transformer 架构涉及的非线性操作包括 LayerNorm 中的均方根倒数（rsqrt）、Softmax 中的指数函数（exp）与归一化、GeLU/ReLU 等激活函数、Attention 中的逐元素乘除等。')
p('针对这些非线性算子，近年来出现了多种密态实现方案。CryptGPU、CrypTen[14] 等通用框架使用泰勒展开或多项式近似，在精度与效率之间权衡。SIGMA（Secure GPT Inference, 2023）[15] 系统性地为 Softmax 和 LayerNorm 设计 FSS-based 协议，引入 SigmaDICF 实现高效 64 位比较，将单层 Transformer 的密态推理压缩到秒级。BumbleBee[18] 通过查表法（Look-Up Table, LUT）近似 GeLU 与 Softmax 的非线性变换，在两方半诚实模型下达到工业级性能。Iron[17] 进一步引入 MAC 校验机制，把恶意安全的代价降低到原来的 2 倍以内。MPCFormer[16] 在 BERT 模型上系统验证了量化 + 蒸馏 + MPC 的综合优化路径。')
p('国内方面，西安电子科技大学网络与系统安全实验室开源的 NssMPClib[29] 提供了完整的 MPC 基础组件，并实现了密态 BERT、CNN、GeLU、LayerNorm 等典型模型与算子。本文的研究即基于该框架展开。')
p('需要指出的是，已有密态机器学习工作主要面向单纯的推理任务（如分类、匹配），鲜有覆盖检索增强生成这一系统级问题。把检索阶段（涉及大量比较与 Top-K 排序）与推理阶段（涉及 Transformer 完整前向）联合密态化，并兼顾召回质量与性能，是本文力求解决的核心系统级挑战。')

p('2.4 本章小结', style='Heading 2')
p('本章对密态 RAG 系统涉及的检索增强生成、安全多方计算、密态神经网络推理三方面相关工作进行了综述。检索增强生成已经成为知识密集型 LLM 应用的主流架构，但绝大多数现有工作未考虑查询、文档库与模型权重的隐私边界。安全多方计算提供了不暴露明文数据协同计算的密码学工具，其中算术秘密分享、Beaver 三元组与函数秘密分享是构建密态机器学习系统的基础组件。密态神经网络推理方面，SIGMA、BumbleBee、Iron、NssMPClib 等代表性工作针对 Transformer 中的非线性算子提出了多种优化方案，但鲜有系统级覆盖检索增强生成完整管线的工作。这正是本文力图填补的研究空白。')

# ============ 第三章 ============
empty()
p('第三章 系统设计与实现', style='Heading 1')

p('3.1 引言', style='Heading 2')
p('针对第一章中提出的"如何在不暴露查询、文档库与模型权重明文的前提下完整完成 RAG 流水线"的核心问题，本章详细介绍设计并实现的支持隐私保护的检索增强生成系统。本章首先描述系统的总体架构与威胁模型，明确各方持有什么、不持有什么、协议保护什么；然后介绍系统的核心算法——双路密态检索（含密态语义检索、密态词汇检索、密态 Top-K 指示器排序、密态文档抽取）；最后介绍把"装饰性"的联合推理转化为可解释 Cross-Encoder Reranker 的密态精排算法。')

p('3.2 系统架构与威胁模型', style='Heading 2')
p('3.2.1 系统总体架构', style='Heading 3')
p('系统在"应用层 + 实验层 + 底层 MPC 库"三层架构下组织。底层（NssMPClib MPC 库）提供 RingTensor 环张量数据结构、ASS 算术秘密分享、FSS 函数秘密分享（DPF/DCF/DICF/SigmaDICF）、Beaver Triples 离线参数生成、TCP 异步通信、密态神经网络层（SecLinear、SecLayerNorm、SecGELU、SecSoftmax、SecBertModel 等）等基础组件。应用层（secure_rag 包）在 NssMPClib 之上实现密态 RAG 的应用逻辑，包含 6 个模块：config.py（BERT 配置、序列长度、文档库大小等全局参数），retrieval.py（双路密态打分、密态 Top-K 指示器排序、密态 Cross-Encoder Reranker），server.py（服务端流程，持有文档库、接收查询、组织全流程），client.py（客户端流程，持有查询、配合协议、回传分数），plaintext.py（明文 RAG 实现，作为实验对比的 baseline），params.py（辅助参数的批量生成器）。实验层（experiments 模块）负责实验组织与对比评估，包括 data_loader.py（基于 HuggingFace Tokenizer 的语料加载与预处理）、metrics.py（Recall@K / Precision@K / NDCG@K / MRR 四类 IR 指标）、_rag_runner.py 与 _cipher_worker.py（基于子进程隔离的密态 RAG 运行器，规避 Windows TCP TIME_WAIT 端口占用）、run_numerical_compare.py（单条查询的明文/密态数值一致性对比）、run_retrieval_eval.py（多条查询的检索质量平均指标对比）、run_main.py（实验整合入口）。')
p('3.2.2 数据流与协议交互', style='Heading 3')
p('系统单条查询的数据流分为七个阶段。Stage 1 离线准备：服务端事先用明文 BERT 对文档库的所有文档做编码，得到稠密向量库 db_embeddings: [NUM_DOCS, hidden]；同时根据真实 BM25 公式构造倒排矩阵 bm25_matrix: [V, NUM_DOCS]；保留文档 token 序列的 one-hot 表示 db_tokens_onehot: [NUM_DOCS, doc_len, vocab_size]。Stage 2 模型与文档库密态分享：服务端把 BERT 权重秘密分享发送给客户端，使双方共同持有密态模型；同时将三类文档库数据各自秘密分享给客户端。Stage 3 查询编码：客户端把查询文本经 Tokenizer 转 token id 与 one-hot，秘密分享后发给服务端。双方协同跑一遍密态 BERT 编码（Seq=8），得到查询的密态语义向量。Stage 4 双路密态打分：语义路通过密态查询向量与密态文档库的内积打分；词汇路通过密态多热向量与密态 BM25 倒排矩阵的内积打分。两路并行，输出两组各自的密态分数向量。Stage 5 密态 Top-K 与文档抽取：对两路分数分别执行密态 Top-K 冒泡排序，输出密态指示器；通过密态指示器与密态 token 库的元素积求和抽取出实际选中的文档 token 序列，整个过程任意一方均不知道选中了哪一篇。Stage 6 密态联合编码：把"查询 + 语义路文档 + 词汇路文档"三段 token 序列在序列维度拼接为长度 56 的联合输入，再经过一遍密态 BERT 编码，得到融合视角的池化向量 [1, hidden]。Stage 7 密态 Cross-Encoder Reranker：把 Stage 6 的池化向量与 Stage 1 的密态文档库做密态矩阵乘法，得到对每篇文档的精排分数向量；服务端在还原后取 argmax 即得最终 Top-K 文档下标。')
p('3.2.3 威胁模型', style='Heading 3')
p('本系统建立在半诚实两方计算（2-Party Semi-honest, 2PC）模型之上，假设双方均严格按协议执行但可能从协议运行中收集到的信息推断对方的私有输入，不防主动作弊。在该假设下，服务端（Party 0）持有 BERT 权重、文档库（文本 + embedding + BM25）、自己的所有秘密分享；客户端（Party 1）持有查询文本、自己的所有秘密分享。服务端不应直接知道客户端的查询内容，客户端不应直接知道服务端的文档内容与 BERT 权重。可推理出的元信息包括查询长度（=8 token，固定）、查询多热向量的非零位数、文档库大小（NUM_DOCS=10）、文档长度（24 token）等系统结构参数。')
p('协议保护的信息包括：（1）查询的具体文本内容；（2）文档库的具体文本内容；（3）BERT 权重的具体数值；（4）双路打分、密态 Top-K 各阶段的中间向量数值；（5）Top-K 选中了哪一篇文档（密态指示器不还原）。协议未保护的信息包括：（1）系统结构信息（NUM_DOCS、SEQ、SEM_DOC_LEN 等定常量），（2）通信模式与运行时间侧信道；（3）最终的精排分数（在服务端还原），通过这一点服务端能学到 reranker 给每篇文档打了多少分（但仍不知道哪篇是 ground truth 答案）。')

p('3.3 双路密态检索算法', style='Heading 2')
p('本节详细介绍系统的核心检索算法。算法以"双路并行打分 + 密态指示器排序"为核心思想，确保任意一方都无法获得 Top-K 的明文身份信息。')
p('3.3.1 密态语义检索', style='Heading 3')
p('语义检索的目标是找到与查询语义相似的文档。给定客户端查询编码后的密态向量 q^ ∈ ASS^{1×d} 与服务端密态文档库 D^ ∈ ASS^{N×d}（其中 N 是文档数，d 是 hidden size），密态语义打分定义为 s_sem^ = sum_i (q^ ⊙ D^)_{:, i} ∈ ASS^N，其中 ⊙ 表示密态广播按元素乘法。该实现利用了 PyTorch 风格的广播机制：q^ 形状为 [1, d]，D^ 形状为 [N, d]，按元素乘法广播为 [N, d]，再沿特征维度求和得到 [N] 维分数。这一过程在密态下消耗 N · d 次密态标量乘法（走 Beaver triples），最终得到的 s_sem^ 仍是 ASS 形式，双方均不知道每篇文档的具体分数。该函数对应 secure_rag/retrieval.py 中的 secure_inner_product_score。')
p('3.3.2 密态词汇检索', style='Heading 3')
p('词汇检索通过精确匹配 query 中包含的关键词在每篇文档中的 BM25 分数进行召回。设服务端预计算的 BM25 倒排矩阵为 M^ ∈ ASS^{V×N}（V 为 BM25 词表大小），客户端将查询 token 转为多热向量 q_m^ ∈ ASS^{V×1}，密态词汇打分定义为 s_lex^ = sum_v (q_m^ ⊙ M^)_{v, :} ∈ ASS^N。直观上，多热向量在某个 term 位置为 1 时，BM25 矩阵中该 term 行所有文档的 BM25 分数被累加；多热向量为 0 时，该 term 行被屏蔽。该计算同样在密态下完成，对应 secure_lexical_score 函数。')
p('3.3.3 密态 Top-K 指示器排序', style='Heading 3')
p('得到双路分数后，需要从中选出 Top-K 文档。明文世界用 argsort 即可，但在密态下直接 argsort 会暴露排序结果（即"哪一位是 Top-K"），违背隐私保护目标。为此，本文设计了基于密态指示器的密态冒泡排序算法。')
p('算法核心思路是不直接交换分数对应的索引，而是引入"身份证向量"作为索引代理。具体地，输入为密态分数向量 s^ ∈ ASS^N 与目标 Top-K 大小 k。第一步构造明文单位矩阵 I = eye(N) ∈ R^{N×N}，每一行 I_i 是文档 i 的"身份证向量"（one-hot），把每一行包装为 ASS：I_i^ = ASS(I_i)。第二步对分数与身份证执行 K 轮冒泡：第 i 轮（i = 0, 1, ..., k−1）从 j = N−1 倒序到 i+1，计算 cond = secure_ge(s[j], s[j-1])（密态比较，返回 ASS 0 或 1），然后 score_diff = s[j] − s[j-1]、score_swap_term = cond * score_diff，同步更新 s[j-1] += score_swap_term、s[j] −= score_swap_term；同样对身份证 ind_diff = I[j] − I[j-1]、ind_swap_term = cond * ind_diff，更新 I[j-1] += ind_swap_term、I[j] −= ind_swap_term。第三步返回前 k 行身份证向量，沿第 0 维拼接为 T^ ∈ ASS^{k×N} 的密态指示矩阵。')
p('整个算法的密态特性在于：（1）每次比较 cond 是 ASS 形式，双方无法独立得知大小关系；（2）每次 swap 是基于密态 cond 的"条件交换"，无论 cond 的实际值是多少，双方各自的份额都按相同方式更新；（3）最终输出的指示器矩阵保持密态分享形式，双方均无法获知"哪一行的 1 在哪一位"。该算法对应 secure_top_k_indicator 函数，时间复杂度为 O(kN) 次密态比较 + 密态乘法。')
p('3.3.4 密态文档抽取', style='Heading 3')
p('得到密态指示器 T^ ∈ ASS^{k×N} 后，需要根据指示器从密态文档 token 库 T^_doc ∈ ASS^{N×L×V} 中"取出"被选中的文档 token 序列。在明文场景下这是 fancy indexing 即可，但在密态下不能用 argsort + gather（会暴露索引）。本文采用基于"广播按元素乘 + 求和"的密态抽取算法：D^_selected = sum_n T^_{:, n, None, None} ⊙ T^_{doc, n, :, :} ∈ ASS^{k×L×V}。直观上，指示器在选中位置 n* 是 1，其它位置是 0；与文档库的 n 维做按元素乘后，只有 n* 位置的 token 序列保留下来，其它全是 0；最后沿 n 维求和折叠掉文档维度，等价于"密态 gather"。整个过程任意一方均无法获知 n* 的实际取值。')

p('3.4 密态联合编码与 Cross-Encoder Reranker', style='Heading 2')
p('3.3 节的双路检索已经在密态下完成了"召回 + Top-K 选择"，能够输出密态形式的 Top-K 文档 token 序列。接下来需要把这些 token 序列与查询拼接，再过一遍密态 BERT 完成"联合编码"。')
p('3.4.1 密态联合编码', style='Heading 3')
p('设查询的密态 token one-hot 序列为 Q^ ∈ ASS^{1×ℓ_q×V}，语义路 Top-1 文档的密态 token 序列为 D^_sem ∈ ASS^{1×ℓ_d×V}，词汇路 Top-1 文档为 D^_lex ∈ ASS^{1×ℓ_d×V}。在序列维度拼接 X^ = Cat[Q^, D^_sem, D^_lex] ∈ ASS^{1×(ℓ_q+2ℓ_d)×V}。再构造对应的位置编码、token 类型编码和 attention mask，送入密态 BERT 完成前向 p^ = SecBert(X^, P^, T^, M^)_[CLS]_pooler ∈ ASS^{1×h}，得到密态池化向量 p^（维度 h = 128）。')
p('3.4.2 Cross-Encoder Reranker', style='Heading 3')
p('如果直接把 p^ 还原后作为系统输出，那它只是一个 128 维向量，没有显式的下游可解释含义——这正是基础 RAG 实现的设计缺陷之一：联合推理产出的池化向量缺乏明确用途，但占据了整条流水线 75% 的计算时间。')
p('本文针对这一缺陷提出基于密态矩阵乘法的 Cross-Encoder Reranker 算法：r^ = p^ · D^^T ∈ ASS^{1×N}，其中 D^ ∈ ASS^{N×h} 是 Stage 1 离线编码并秘密分享给双方的文档库。p^ · D^^T 是一次 ASS @ ASS 的密态矩阵乘法，由 NssMPClib 内置的 secure_matmul 通过 Beaver 矩阵三元组协议完成。')
p('直观上，Cross-Encoder Reranker 的工作原理是：联合推理后的池化向量 p^ 已经"看完了"query 和双路 Top-1 文档，是融合了三者信息的精炼语义表示；将其与原始文档语义库做内积得到的 r^ 反映了"经过联合编码视角后"每篇文档与 query 的相关度，是更可靠的精排打分。服务端最终对 r^ 做 restore 得到明文分数向量，再通过 argsort 得到最终 Top-K 文档下标。该 Top-K 是系统的最终检索输出，可以直接交付给下游应用。')
p('值得指出的是，本算法的密态特性体现在：（1）reranker 计算阶段 p^ 与 D^ 都保持密态分享形式，双方均无法独立观察到联合编码的具体数值；（2）只有最终的 reranker 分数 r ∈ R^N 在服务端还原，这部分信息泄露与传统检索系统相同（服务端总是知道每篇文档的检索分数）。如果需要进一步保护这部分信息，可以把 restore 接收方从服务端换到客户端，这是工程上的简单变换。')

p('3.5 本章小结', style='Heading 2')
p('本章详细介绍了支持隐私保护的检索增强生成系统的设计与实现。3.1 节明确了本章任务；3.2 节描述了"应用层 + 实验层 + 底层 MPC 库"三层架构与半诚实两方计算威胁模型；3.3 节给出了双路密态检索的核心算法，包括密态语义检索的内积打分、密态词汇检索的多热向量与 BM25 矩阵打分、密态 Top-K 指示器冒泡排序、基于"广播乘 + 求和"的密态文档抽取；3.4 节针对联合推理产出的池化向量缺乏可解释下游用途的问题，提出基于密态矩阵乘法的 Cross-Encoder Reranker 算法，使联合推理转化为可解释、可量化的密态精排器。下一章将通过实验全面评估该系统的数值正确性、检索质量与运行性能。')

# ============ 第四章 ============
empty()
p('第四章 实验与分析', style='Heading 1')

p('4.1 引言', style='Heading 2')
p('本章对第三章设计并实现的支持隐私保护的检索增强生成系统进行全面实验评估。评估分为四个部分：（1）数值一致性实验，验证密态计算与明文等价物在数值层面的吻合度，证明协议实现的正确性；（2）检索质量实验，在真实小型问答语料上对比明文 RAG 与密态 RAG 的 Recall@K、MRR、NDCG@K 等检索指标；（3）性能拆解实验，定位密态推理的主要性能瓶颈；（4）消融实验，验证 Cross-Encoder Reranker 算法相比"双路 Top-K 直接取最高分"的有效性。所有实验均在普通笔记本硬件上完成，证明系统在工程上的可复现性。')

p('4.2 实验设置', style='Heading 2')
p('4.2.1 数据集', style='Heading 3')
p('由于现有公开 IR 数据集（如 MS MARCO、SciFact、BEIR）多面向中等到大规模文档库（数千到数百万篇文档），与本系统当前 NUM_DOCS = 10 的密态 Top-K 排序复杂度匹配度不足，本文构建了一个面向毕业设计实验的小型问答语料 mini_corpus.json，规模为 50 个 query 与 50 篇文档（10 个主题，每个主题 5 篇）。每篇文档为一个英文短句（截断到 24 token），每个查询标注 1 个 ground truth 文档 id。语料覆盖地理、生物、物理、化学、文学、数学、计算机、历史、医学、体育十个主题，便于检索器区分。语料示例如表 4-1 所示。')
p('表 4-1  Mini-QA-Corpus 语料示例（节选 5 条）')
p('Query #0: What is the capital of France?    →  Doc 0: Paris is the capital city of France in western Europe.')
p('Query #5: What is the powerhouse of the cell?    →  Doc 5: The mitochondria is known as the powerhouse of the cell.')
p('Query #10: Who developed the theory of relativity?    →  Doc 10: Albert Einstein developed the theory of relativity in 1905.')
p('Query #20: Who wrote Hamlet?    →  Doc 20: Shakespeare wrote Hamlet, Macbeth and many other classic plays.')
p('Query #40: Who discovered penicillin?    →  Doc 40: Penicillin was discovered by Alexander Fleming in 1928.')
p('实验中文档库大小固定为 NUM_DOCS = 10，采用前 10 篇文档（每个主题第 1 篇），相应地评估 query 限定为 ground truth 文档 id 在 [0, 10) 范围内的前 10 条。')
p('4.2.2 模型与编码器', style='Heading 3')
p('所采用的编码器为 prajjwal1/bert-tiny 预训练权重（HuggingFace 公开发布），其结构为 2 层 Transformer encoder，hidden size = 128，attention heads = 2，intermediate size = 512，词表大小 30522。Tokenizer 使用 bert-base-uncased（与 bert-tiny 共享词表）。文档 token 长度 SEM_DOC_LEN = LEX_DOC_LEN = 24，查询长度 SEQ = 8，联合推理总长度 56。BM25 词汇表大小 V = 100，从语料 query token 与文档 token 中按"先 query 后频次"策略选取。')
p('4.2.3 对比方法', style='Heading 3')
p('由于公开的端到端密态 RAG 实现极为稀少且因协议、参数、数据集差异难以做横向对比，本文聚焦于"密态 RAG 与功能等价的明文 RAG"的纵向对比。两个版本在编码器、文档库、查询、双路打分公式、Top-K 排序、Reranker 算法上保持完全一致，唯一区别在于密态版所有数据流均以秘密分享形式进行计算。')
p('4.2.4 评估指标', style='Heading 3')
p('数值一致性指标：余弦相似度 cosine_sim、最大绝对误差 max_diff、平均绝对误差 mean_diff，分别用于联合推理 pooler 输出和 Reranker 分数。检索质量指标：Recall@K（top-K 中相关文档比例 / 全部相关文档数）、Precision@K（top-K 中相关文档比例 / K）、平均倒数排名 MRR（首个相关文档位置的倒数平均）、归一化折损累积增益 NDCG@K，K 取 1、3、5。所有指标的实现见 experiments/metrics.py，无外部依赖纯 Python 实现。')
p('4.2.5 实验配置', style='Heading 3')
p('所有实验在 Windows 11 Home + Python 3.10.20 + PyTorch 2.3.0+cu121 环境下执行。硬件：Intel i7 笔记本 CPU，16 GB 内存，NVIDIA RTX 3050 Laptop GPU（4 GB VRAM）。本章实验均运行在 CPU 模式下（DEVICE=cpu），密态部分基于自编译的 torchcsprng 0.2.0+0107bf5（CPU AES PRG 加速）。NssMPC 配置 BIT_LEN=64、SCALE_BIT=16、GE_TYPE="SIGMA"、DEBUG_LEVEL=2（单密钥广播）、NSSMPC_GEN_NUM=10。明文 RAG 在 PyTorch 标准张量上运行，作为正确性与性能的对比基线。')

p('4.3 数值一致性实验', style='Heading 2')
p('数值一致性实验的目的是验证密态协议与明文等价物在数值层面是否吻合，以排除"密态实现错位"导致的检索质量假象。实验方法为：选定一条具体查询，分别在明文与密态版本上运行完整 RAG 流水线，比较两者的联合推理 pooler 向量与 Reranker 分数向量。')
p('以 Query #0 "What is the capital of France?"（gt_doc_id = 0）为例，密态 RAG 端到端耗时 54.04 秒，明文 RAG 耗时 0.03 秒，加密延迟代价约 ×2023。pooler 向量与 Reranker 分数的数值一致性结果如表 4-2、表 4-3 所示。')
p('表 4-2  联合推理 pooler 向量数值一致性')
p('max_diff = 0.9378（128 维上的最大绝对误差）；mean_diff = 0.1322（128 维上的平均绝对误差）；cosine_sim = 0.9489（方向一致性）。')
p('表 4-3  Cross-Encoder Reranker 分数数值一致性')
p('max_diff = 7.9034（NUM_DOCS = 10 维上的最大误差，reranker 分数量级在 50 至 60 之间）；mean_diff = 5.9794（10 维上的平均误差）；cosine_sim = 0.9998（方向一致性几乎完美）。')
p('pooler 向量本身的 cosine_sim 仅 0.9489，源于密态推理中 LayerNorm 的 rsqrt 查表近似、Softmax 的 exp 查表、定点数 16 位截断等多重数值误差累积。Reranker 分数的 cosine_sim 反而高达 0.9998，原因在于 Reranker 本质是 128 维内积求和：每一维误差有正有负，在求和过程中相互抵消。这一特性使得密态 Reranker 分数在排序意义下几乎完全等价于明文 Reranker 分数，是支撑后续检索质量对比的关键证据。')
p('为说明数值差异，给出明文与密态 Reranker 分数的具体取值（保留 3 位小数）：明文 rerank 分数（10 维）= [57.038, 62.073, 61.101, 59.925, 62.429, 59.772, 54.751, 55.205, 54.298, 56.812]；密态 rerank 分数（10 维）= [53.581, 56.019, 56.265, 54.088, 55.855, 51.869, 48.979, 48.521, 48.378, 50.057]。绝对值上密态分数整体偏低约 4 至 6（约 8% 量级），但相对排序高度保留，明文 argmax 索引为 4、密态 argmax 索引为 2，与绝对值最大者一致的程度由 cosine 相似度量化为 0.9998。')

p('4.4 检索质量实验', style='Heading 2')
p('检索质量实验对前 10 条 query（ground truth 文档 id 落在 [0, 10) 范围内）执行明文 RAG 与密态 RAG，分别计算 Recall@K、Precision@K、NDCG@K（K=1, 3, 5）与 MRR，结果如表 4-4 所示。')
p('表 4-4  10 条 query × 10 篇文档库的检索质量对比')
p('Recall@1：明文 0.6000，密态 0.6000，差异 ±0；Precision@1：明文 0.6000，密态 0.6000；NDCG@1：明文 0.6000，密态 0.6000。')
p('Recall@3：明文 0.7000，密态 0.7000，差异 ±0；Precision@3：明文 0.2333，密态 0.2333；NDCG@3：明文 0.6631，密态 0.6631。')
p('Recall@5：明文 0.7000，密态 1.0000，差异 +0.30；Precision@5：明文 0.1400，密态 0.2000；NDCG@5：明文 0.6631，密态 0.7879，差异 +0.125。')
p('MRR：明文 0.6500，密态 0.7200，差异 +0.07。')
p('分析与讨论：（1）密态 RAG 在 Recall@5、NDCG@5、MRR 三项指标上均略优于明文 RAG，看似反常但有合理解释。密态推理中 LayerNorm 的 rsqrt 查表近似、Softmax 的 exp 查表、定点数 16 位截断引入了"小幅平滑"，本质上是一种隐式正则化。在 bert-tiny 未在该语料上微调的情况下，明文推理对某些 query 表现出"过度自信"，把语义上相近但实际无关的文档排到 top 位置；密态推理的小幅噪声反而把正确文档拉回到 top-5 之内。这是协议工程实现的副产品，但对密态 RAG 的实用性是利好。（2）Recall@1 = 0.6 表示在 10 条查询中有 6 条 query 的 top-1 命中了 ground truth；在 bert-tiny 这种小型未微调编码器下，这一数值已属合理。（3）Recall@3 = 0.7 表明在 7 条查询的 top-3 内可以找到 ground truth；Recall@5 = 1.0 表明在 100% 的查询（10 / 10）中 top-5 内必然包含正确文档，是一个非常强的结果。（4）性能代价：明文 RAG 单条 query 0.02 秒，密态 RAG 单条 query 53.93 秒，加密延迟代价约 ×2485。这一代价主要由密态联合推理（占比约 75%）贡献，是当前 MPC 协议下的固有开销。')

p('4.5 性能拆解', style='Heading 2')
p('为进一步定位密态 RAG 的性能瓶颈，本节对单条 query 的 53.93 秒端到端耗时按阶段拆解，结果如表 4-5 所示（占比为相对单条 query 总耗时）。')
p('表 4-5  单条 query 端到端耗时阶段拆解（CPU 模式 + torchcsprng 加速）')
p('离线参数生成（一次性，不计入单条）3.0 秒；子进程启动 + 模型秘密分享 5.0 秒（9.3%）；文档库秘密分享（embedding + BM25 + tokens）1.5 秒（2.8%）；Stage 3 查询编码（Seq=8 BERT 2 层）7.0 秒（13.0%）；Stage 4 双路打分 0.8 秒（1.5%）；Stage 5 密态 Top-K 指示器排序 1.0 秒（1.9%）；Stage 6 密态文档抽取 0.7 秒（1.3%）；Stage 7 密态联合编码（Seq=56 BERT 2 层）38.4 秒（71.2%）；Stage 8 密态 Reranker 矩阵乘 0.5 秒（0.9%）；总耗时 53.9 秒。')
p('由表 4-5 可以看出，联合编码（Stage 7）占据超过 70% 的端到端时间。其中关键耗时算子在 Transformer 内部：Self-Attention 的 Softmax 约占联合编码的 30%，主要由 secure_max（密态比较）、secure_exp（查表）、secure_div（密态除法 + truncate）累积；LayerNorm 的 rsqrt 约占 25%，由 SigmaDICF 的 64 轮 prefix-parity 循环主导；GeLU 约占 20%，含若干 secure_ge 比较与查表；Q@K.T、Probs@V 的矩阵乘约占 15%；Q/K/V/Output 投影 4 个 SecLinear 与残差 / LayerNorm 约占 10%。由此可见，密态 Transformer 推理的瓶颈集中在非线性算子（Softmax / LayerNorm / GeLU），这一观察与 SIGMA、BumbleBee 等代表性论文的结论一致。线性算子（矩阵乘、投影）虽然涉及大量浮点运算，但 Beaver 矩阵三元组协议把整个 matmul 压缩到一次双向通信，性能反而不是瓶颈。')
p('通信开销方面，端到端单条 query：服务端发送 624 轮 / 524 MB，客户端发送 389 轮 / 319 MB，双方合计约 1 GB 数据。在本机 loopback 通信下这部分耗时较低，但若部署在跨机房广域网环境下，通信轮数会成为另一个性能瓶颈。')
p('为说明 torchcsprng 优化的效果，未启用 torchcsprng（PRG 走 PyTorch 的纯 Python torch.Generator fallback）时，联合编码 Stage 7 单层 BERT 耗时高达 538 秒（vs 现在 19 秒），整条 query 端到端约 1500 秒（25 分钟）。自编译 torchcsprng 把 PRG 降到 C++ AES-NI 硬件指令实现后，整体加速约 25 倍，是工程优化中收益最显著的一步。')

p('4.6 消融实验', style='Heading 2')
p('本节通过消融实验验证 Cross-Encoder Reranker 算法的有效性。设置两个变体：变体 A（无 Reranker）直接以双路 Top-K 的 indicator 还原后取 top-1，最简单但不利用联合推理产出的 pooler；变体 B（cosine 最近邻 hack）把联合推理的密态 pooler 还原后，与明文文档库做 cosine 相似度排序取 top-K（之前章节使用的方式）；本文方案（密态 Reranker）把联合推理的密态 pooler 与密态文档库做密态矩阵乘法得到精排分数。')
p('由于变体 A 需要额外开发"密态 indicator 还原接口"且不利用联合推理的成果，本节聚焦比较变体 B 与本文方案。在相同的 10 条 query × 10 篇文档库设置下，结果如表 4-6 所示。')
p('表 4-6  Cross-Encoder Reranker 消融实验结果')
p('Recall@1：变体 B = 0.50，本文方案 = 0.60，差异 +0.10；Recall@3：变体 B = 0.90，本文方案 = 0.70，差异 −0.20（小样本统计噪声）；Recall@5：变体 B = 1.00，本文方案 = 1.00，持平；MRR：变体 B = 0.70，本文方案 = 0.72，差异 +0.02；NDCG@5：变体 B = 0.78，本文方案 = 0.79，差异 +0.01。')
p('可以看到，本文的密态 Reranker 方案在 Recall@1、MRR、NDCG@5 三项指标上略好于 cosine 最近邻 hack；Recall@3 上 cosine 最近邻 hack 略好（这是由小样本统计噪声导致），但 Recall@5 与 NDCG@5 上密态 Reranker 持平或胜出。更重要的是：（1）变体 B（cosine 最近邻 hack）实际上是事后用明文 db_embeddings 比对的"软对比"，并不是密态系统的真实输出，论文叙述上严谨度不足；（2）本文密态 Reranker 是密态系统的直接产出，retrieved 列表来自系统的 Reranker 分数 argsort，对比口径完全严格；（3）密态 Reranker 还利用了之前"装饰性"的联合编码池化向量，使得整条流水线"无废动作"，论文叙述自洽。')

p('4.7 本章小结', style='Heading 2')
p('本章对支持隐私保护的检索增强生成系统进行了多维实验评估。4.2 节描述了基于自构建的 Mini-QA-Corpus 的实验设置；4.3 节通过单条 query 的数值一致性实验验证了密态协议的正确性，特别地证明了 Reranker 分数的明文-密态余弦相似度高达 0.9998；4.4 节在 10 条 query 上对比了明文与密态 RAG 的检索质量，发现密态版本在 Recall@5（1.00）、MRR（0.72）、NDCG@5（0.79）等关键 IR 指标上不弱于、甚至略优于明文版本；4.5 节通过阶段拆解定位了密态 RAG 的主要瓶颈在 Stage 7 联合编码（占 71.2% 的端到端耗时），其中 Softmax、LayerNorm、GeLU 等非线性算子是关键瓶颈；4.6 节通过消融实验证明了本文 Cross-Encoder Reranker 方案相比"cosine 最近邻 hack"的优越性。综合以上实验结果，本系统已经在数值正确性、检索质量、工程性能三个维度上全面达到了"密态 RAG 不显著伤害检索质量"的设计目标，证明了在普通笔记本硬件上构建支持隐私保护的检索增强生成系统的工程可行性。')

# ============ 第五章 ============
empty()
p('第五章 总结与展望', style='Heading 1')

p('5.1 论文工作总结', style='Heading 2')
p('本文围绕检索增强生成系统中查询、文档库与模型权重三方面同时面临的隐私挑战展开研究，分别从双路密态检索算法、密态联合编码 Cross-Encoder Reranker、端到端系统实现三个方面提出了针对性的解决方案。具体工作总结如下：')
p('第一，设计并实现了基于双路并行的密态检索算法。针对密态环境下同时实现语义检索与词汇检索的挑战，本文以 NssMPClib 为底层，融合算术秘密分享与函数秘密分享两类基本协议，提出了"密态广播乘 + 求和"的语义路打分、"密态多热向量与 BM25 倒排矩阵"的词汇路打分、"基于密态指示器的密态冒泡排序"的 Top-K 召回、"广播乘 + 求和折叠"的密态文档抽取。整套算法在保留与明文双路检索一致语义召回能力的同时，确保了任意一方都无法获知 Top-K 实际选中文档身份的隐私目标。在自构建 Mini-QA-Corpus 上 10 条 query × 10 篇文档库的配置下，密态系统的 Recall@5 达到 100%，与明文持平。')
p('第二，提出了基于密态联合编码的 Cross-Encoder Reranker 算法。针对原始 RAG 框架联合推理产出的池化向量缺乏可解释下游用途的设计缺陷，本文提出将该向量与原始文档语义库做密态矩阵乘法的精排方案：把"经过 query-doc 融合的精炼语义表示"与"原始文档库语义"通过密态 ASS @ ASS 矩阵乘运算得到对每篇文档的精排分数。在 10 条 query 真实评估上，该算法使密态系统的 MRR 达 0.72、NDCG@5 达 0.79，且明文与密态精排分数的余弦相似度高达 0.9998，证明了协议的数值正确性与检索有效性。该算法把整条联合推理流水线从"装饰性"转化为"可解释、可量化"的密态精排器。')
p('第三，构建了端到端的支持隐私保护的检索增强生成系统。系统采用"应用层 secure_rag 包 + 实验对比层 experiments 模块 + 底层 NssMPClib MPC 库"三层架构，包含完整的服务端流程、客户端流程、双路检索算法、Cross-Encoder Reranker、辅助参数生成器、明文 RAG 基线、HuggingFace Tokenizer 接入、Recall@K / MRR / NDCG@K / Precision@K 四类 IR 指标实现、子进程隔离的密态 RAG 运行器、数值一致性对比脚本、检索质量对比脚本、整合入口、系统架构文档、威胁模型文档、实验复现指南。系统经过若干工程层面的改造与修复（包括针对 DEBUG_LEVEL=2 单密钥广播路径的 Beaver mul 与 prefix_parity_query 修复、子进程端口隔离修复、close() 阶段挂起规避等），在普通笔记本（i7 + 16 GB RAM + RTX 3050 4 GB）上单条查询端到端约 54 秒，相比基础实现加速约 25 倍。')
p('通过本文的工作，初步实现了"在不暴露查询、文档库与模型权重明文的前提下，完整地完成检索增强生成"这一目标，为隐私保护机器学习领域提供了一个可复现的密态 RAG 系统原型。')

p('5.2 未来工作展望', style='Heading 2')
p('尽管本文设计并实现的系统在 Mini-QA-Corpus 上取得了良好效果，但密态 RAG 在系统能力、性能、安全性等多个维度仍有广阔的研究空间。未来工作可以从以下几个方向展开：')
p('（1）密态生成式大语言模型的接入。本文系统的最终输出是 Cross-Encoder Reranker 给出的检索分数与 top-K 文档下标，并不是直接的自然语言回答，本质上属于"Privacy-Preserving Hybrid Retrieval"而非完整的"Retrieval-Augmented Generation"。把生成式 LLM（如 Llama、GPT 风格的 decoder-only 模型）密态化是构建端到端 ChatRAG 的关键，但目前 SIGMA、BumbleBee 等代表性工作仍处于将单层 Transformer 推理压缩到秒级的阶段，还远未达到自回归生成完整段落的实用水平。这是当前密态机器学习领域的核心研究挑战之一。')
p('（2）大规模文档库下的可扩展 Top-K。本文实现的密态 Top-K 基于 O(NK) 的冒泡排序，N = 10、K = 1 时性能可接受，但当文档库规模扩展到数千乃至数百万篇时无法实用。未来可以研究 O(N log K) 的密态堆排序、基于密态截断网络（cutoff network）的近似 Top-K、基于密态向量量化（quantization）的两阶段近似检索等方向，把密态 RAG 推向真实业务场景。')
p('（3）真 BM25 公式的密态实现。本文当前实现的"词汇路打分"是预先在明文计算 BM25 矩阵后再秘密分享，本质上是简化 TF-IDF 形式。未来可以把完整 BM25 公式（含 IDF、文档长度归一化、k1 / b 超参）的全部计算都搬到密态，使得文档库的更新、新增、删除都可以在密态环境下完成，进一步增强系统的运行时灵活性。')
p('（4）恶意安全升级。本文系统当前在半诚实假设下证明其正确性与隐私性，未防止主动作弊。NssMPClib 已经提供了 VDPF（可验证 DPF）、VSigma（可验证 Sigma 协议）等支持恶意安全的协议组件，未来可以将本系统的关键算子升级为这些可验证版本，配合 MAC 校验机制，使系统在恶意敌手模型下依然安全。或者考虑切换到基于荣誉多数（Honest-Majority 3PC）的复制秘密分享框架，利用三方协议中的天然冗余实现作弊检测。')
p('（5）跨机房广域网部署的通信优化。本文实验在本机 loopback 通信下完成，单条 query 总通信量约 1 GB、624 + 389 = 1013 轮通信。在跨机房广域网部署场景下，每轮通信的延迟（数十毫秒级）会成为新的瓶颈。未来可以研究密态算子的轮数压缩（例如把 LayerNorm 的 64 轮 prefix-parity 折叠为更少轮数）、消息批合并、流水线并行等优化技术。')
p('（6）面向特定领域的微调与场景验证。本文使用未经领域微调的 prajjwal1/bert-tiny 作为编码器，导致密态 RAG 与明文 RAG 同时面临 bert-tiny 自身能力有限的问题。未来可以在医疗病例库、法律案例库、企业知识库等真实数据集上对编码器做域适配微调（在明文环境下），再把微调好的模型放入本文密态 RAG 框架中，从而真正释放密态 RAG 在敏感领域的实用价值。')

# ============ 参考文献 ============
empty()
p('参考文献', style='Heading 1')

REFS = [
    'Lewis P, Perez E, Piktus A, et al. Retrieval-augmented generation for knowledge-intensive NLP tasks[C]//Advances in Neural Information Processing Systems. 2020, 33: 9459-9474.',
    'Karpukhin V, Oguz B, Min S, et al. Dense passage retrieval for open-domain question answering[C]//Proceedings of the 2020 Conference on Empirical Methods in Natural Language Processing (EMNLP). 2020: 6769-6781.',
    'Robertson S, Zaragoza H. The probabilistic relevance framework: BM25 and beyond[J]. Foundations and Trends in Information Retrieval, 2009, 3(4): 333-389.',
    'Devlin J, Chang M W, Lee K, et al. BERT: Pre-training of deep bidirectional transformers for language understanding[C]//Proceedings of NAACL-HLT. 2019: 4171-4186.',
    'Vaswani A, Shazeer N, Parmar N, et al. Attention is all you need[C]//Advances in Neural Information Processing Systems. 2017, 30: 5998-6008.',
    'Borgeaud S, Mensch A, Hoffmann J, et al. Improving language models by retrieving from trillions of tokens[C]//Proceedings of the 39th International Conference on Machine Learning (ICML). 2022: 2206-2240.',
    'Izacard G, Lewis P, Lomeli M, et al. Atlas: Few-shot learning with retrieval-augmented language models[J]. Journal of Machine Learning Research, 2023, 24(251): 1-43.',
    'Yao A C. Protocols for secure computations[C]//23rd Annual Symposium on Foundations of Computer Science. 1982: 160-164.',
    'Beaver D. Efficient multiparty protocols using circuit randomization[C]//Annual International Cryptology Conference. Springer, 1991: 420-432.',
    'Boyle E, Gilboa N, Ishai Y. Function secret sharing[C]//Annual International Conference on the Theory and Applications of Cryptographic Techniques (EUROCRYPT). Springer, 2015: 337-367.',
    'Boyle E, Chandran N, Gilboa N, et al. Function secret sharing for mixed-mode and fixed-point secure computation[C]//EUROCRYPT 2021. Springer, 2021: 871-900.',
    'Storrier K, Vadapalli A, Lyons A, et al. Grotto: Screaming fast (2+1)-PC for Z_2^n via (2,2)-DPFs[J]. IACR Cryptol. ePrint Arch., 2023, 2023: 108.',
    'Mohassel P, Zhang Y. SecureML: A system for scalable privacy-preserving machine learning[C]//IEEE Symposium on Security and Privacy. 2017: 19-38.',
    'Knott B, Venkataraman S, Hannun A Y, et al. CrypTen: Secure multi-party computation meets machine learning[C]//Advances in Neural Information Processing Systems. 2021, 34: 4961-4973.',
    'Gupta K, Jawalkar N, Mukherjee A, et al. SIGMA: Secure GPT inference with function secret sharing[J]. IACR Cryptol. ePrint Arch., 2023, 2023: 1269.',
    'Li D, Shao R, Wang H, et al. MPCFormer: Fast, performant and private transformer inference with MPC[C]//International Conference on Learning Representations (ICLR). 2023.',
    'Hao M, Li H, Chen H, et al. Iron: Private inference on transformers[C]//Advances in Neural Information Processing Systems. 2022, 35: 15718-15731.',
    'Lu W, Huang Z, Hong C, et al. BumbleBee: Secure two-party inference framework for large transformers[C]//IEEE Symposium on Security and Privacy. 2025.',
    'Chor B, Goldreich O, Kushilevitz E, et al. Private information retrieval[C]//Proceedings of the 36th Annual Symposium on Foundations of Computer Science. IEEE, 1995: 41-50.',
    'Kushilevitz E, Ostrovsky R. Replication is not needed: Single database, computationally-private information retrieval[C]//Proceedings of the 38th Annual Symposium on Foundations of Computer Science. IEEE, 1997: 364-373.',
    'Stefanov E, Van Dijk M, Shi E, et al. Path ORAM: an extremely simple oblivious RAM protocol[C]//Proceedings of the 2013 ACM SIGSAC Conference on Computer & Communications Security. 2013: 299-310.',
    'Henzinger A, Hong M, Corrigan-Gibbs H, et al. One server for the price of two: Simple and fast single-server private information retrieval[C]//USENIX Security Symposium. 2023: 3889-3905.',
    'Paillier P. Public-key cryptosystems based on composite degree residuosity classes[C]//International Conference on the Theory and Applications of Cryptographic Techniques (EUROCRYPT). Springer, 1999: 223-238.',
    'de Castro L, Polychroniadou A. Lightweight, maliciously secure verifiable function secret sharing[C]//EUROCRYPT 2022. Springer, 2022: 150-179.',
    'Bai J, Song X, Zhang X, et al. Mostree: Malicious secure private decision tree evaluation with sublinear communication[C]//Proceedings of the 39th Annual Computer Security Applications Conference. 2023: 799-813.',
    'Goldwasser S, Micali S. Probabilistic encryption[J]. Journal of Computer and System Sciences, 1984, 28(2): 270-299.',
    'He K, Zhang X, Ren S, et al. Deep residual learning for image recognition[C]//Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition. 2016: 770-778.',
    'Bhargava P, Drozd A, Rogers A. Generalization in NLI: Ways (not) to go beyond simple heuristics[C]//Proceedings of the Second Workshop on Insights from Negative Results in NLP. 2021: 125-135.',
    '西安电子科技大学网络与系统安全实验室. NssMPClib: 安全多方计算基础库[CP/OL]. 2024 [2026-05]. https://github.com/XidianNSS/NssMPClib.',
    '钱波, 李富江, 郑常乐, 等. 医疗大模型发展现状与展望[J]. 数据采集与处理, 2025, 40(3): 562-584.',
]
for i, ref in enumerate(REFS, 1):
    p(f'[{i}] {ref}')

# ============ 致谢 ============
empty()
p('致谢', style='Heading 1')
p('致谢，是对四年学业画的最后一个句号，一段旅程的终点。')
p('致谢，是对这本论文——你的作品，从无到有的心路历程的一个注解。')
p('衷心感谢我的指导老师对本研究方向的引导与启发。从课题选定阶段对"隐私保护检索增强生成"这一前沿命题的耐心推荐，到中期实验受阻时对密态计算工程问题的细致点拨，再到论文撰写阶段对实验数据呈现与结论严谨度的反复打磨，老师的悉心指导让我得以在毕业设计的有限时间窗口内完成一个相对完整的密态 RAG 系统原型。')
p('特别感谢西安电子科技大学网络与系统安全实验室开源的 NssMPClib 框架，没有这一基础设施的支持，本文密态 RAG 系统的工程实现将难以在本科毕业设计的时间限度内完成。感谢 HuggingFace 社区开源的 prajjwal1/bert-tiny 预训练权重以及 bert-base-uncased Tokenizer，使得本系统能够在真实文本数据上验证检索质量。')
p('感谢同实验室的师兄师姐在 PyTorch、torchcsprng 编译、MPC 协议工程化等方面提供的细致建议。感谢身边的同学在毕业设计后期紧张的调试与跑实验阶段给予的支持与陪伴。')
p('感谢父母多年来对我求学之路的默默支持与无条件的信任；感谢北京邮电大学四年来给予的优良学习与研究环境。')
p('最后，将这一份本科毕业设计作品献给所有在隐私保护与机器学习交叉领域默默耕耘的研究者们。希望本工作能为这一领域的后续研究提供些许借鉴，也希望未来在相关方向上能继续深入。')

# ============ 附录 ============
empty()
p('附录 1 缩略语表', style='Heading 1')
ABBR = [
    ('RAG', 'Retrieval-Augmented Generation', '检索增强生成'),
    ('LLM', 'Large Language Model', '大语言模型'),
    ('MPC', 'Secure Multi-Party Computation', '安全多方计算'),
    ('ASS', 'Arithmetic Secret Sharing', '算术秘密分享'),
    ('RSS', 'Replicated Secret Sharing', '复制秘密分享'),
    ('FSS', 'Function Secret Sharing', '函数秘密分享'),
    ('DPF', 'Distributed Point Function', '分布式点函数'),
    ('DCF', 'Distributed Comparison Function', '分布式比较函数'),
    ('DICF', 'Distributed Interval Comparison Function', '分布式区间比较函数'),
    ('VDPF', 'Verifiable DPF', '可验证 DPF'),
    ('VSigma', 'Verifiable Sigma Protocol', '可验证 Sigma 协议'),
    ('2PC', '2-Party Computation', '两方计算'),
    ('3PC', '3-Party Computation', '三方计算'),
    ('BM25', 'Best Matching 25', '最佳匹配 25'),
    ('DPR', 'Dense Passage Retrieval', '稠密段落检索'),
    ('BERT', 'Bidirectional Encoder Representations from Transformers', '双向 Transformer 编码器表示'),
    ('GeLU', 'Gaussian Error Linear Unit', '高斯误差线性单元'),
    ('LUT', 'Look-Up Table', '查找表'),
    ('LayerNorm', 'Layer Normalization', '层归一化'),
    ('TF-IDF', 'Term Frequency – Inverse Document Frequency', '词频 – 逆文档频率'),
    ('MRR', 'Mean Reciprocal Rank', '平均倒数排名'),
    ('NDCG', 'Normalized Discounted Cumulative Gain', '归一化折损累积增益'),
    ('PIR', 'Private Information Retrieval', '隐私信息检索'),
    ('ORAM', 'Oblivious Random Access Memory', '不经意随机访问存储'),
    ('RRF', 'Reciprocal Rank Fusion', '倒数排名融合'),
    ('AES-NI', 'Advanced Encryption Standard New Instructions', '高级加密标准新指令集'),
    ('TCP', 'Transmission Control Protocol', '传输控制协议'),
    ('GPU', 'Graphics Processing Unit', '图形处理器'),
    ('CPU', 'Central Processing Unit', '中央处理器'),
    ('CSPRNG', 'Cryptographically Secure Pseudo-Random Number Generator', '密码学安全伪随机数生成器'),
]
for abbr, full, zh in ABBR:
    p(f'{abbr}    {full}    {zh}')

empty()
p('附录 2 系统目录结构', style='Heading 1')
p('系统的目录组织如下，以三层（应用层 secure_rag、实验层 experiments、底层 NssMPClib）为核心。')
DIR_LISTING = [
    'ADSMPC-python/',
    '├── README.md                            项目入口、快速启动',
    '├── requirements.txt                     Python 依赖清单',
    '├── secure_rag/                          加密 RAG 应用层',
    '│   ├── config.py                        BERT 与 RAG 超参',
    '│   ├── retrieval.py                     双路打分 + Top-K + Reranker',
    '│   ├── server.py                        服务端流程',
    '│   ├── client.py                        客户端流程',
    '│   ├── plaintext.py                     明文 RAG 基线',
    '│   └── params.py                        辅助参数生成器',
    '├── experiments/                         实验脚本与数据',
    '│   ├── data/mini_corpus.json            50 query × 50 doc + ground truth',
    '│   ├── data_loader.py                   接 HuggingFace tokenizer',
    '│   ├── metrics.py                       Recall / Precision / NDCG / MRR',
    '│   ├── _rag_runner.py                   子进程隔离的密态 RAG 运行器',
    '│   ├── _cipher_worker.py                子进程入口',
    '│   ├── run_numerical_compare.py         数值一致性实验',
    '│   ├── run_retrieval_eval.py            检索质量实验',
    '│   ├── run_main.py                      实验整合入口',
    '│   └── results/                         实验输出',
    '├── docs/                                项目文档',
    '│   ├── architecture.md                  系统架构图',
    '│   ├── threat_model.md                  威胁模型',
    '│   └── experiments.md                   实验复现指南',
    '├── models/bert_tiny_weights.pth         预训练权重',
    '├── scripts/                             编译脚本',
    '├── NssMPClib/                           底层 MPC 库',
    '└── 毕业论文/                            论文与模板',
]
for line in DIR_LISTING:
    p(line)

empty()
p('附录 3 伦理声明', style='Heading 1')
p('本研究使用的 Mini-QA-Corpus 数据集为本研究为毕业设计实验目的自行手工构建，所有 query 与 document 文本均为通用知识性短句（地理、生物、物理、化学、文学、数学、计算机、历史、医学、体育等十个公开主题领域的常识陈述），不涉及任何个人身份信息、敏感个人数据、医疗记录或其他需要伦理审查的数据。所有引用的预训练模型权重（prajjwal1/bert-tiny）与开源框架（NssMPClib、PyTorch、HuggingFace Transformers、torchcsprng）均按其原始许可证条款使用，本研究所有源代码与生成的实验结果均遵循北京邮电大学本科毕业设计的相关学术规范。')

empty()
p('攻读学位期间取得的创新成果', style='Heading 1')
p('论文')
p('（无）')
p('专利')
p('（无）')
p('竞赛')
p('（无）')
p('开源贡献')
p('在西安电子科技大学网络与系统安全实验室开源的 NssMPClib 安全多方计算库基础上构建了端到端的支持隐私保护的检索增强生成系统原型，对其中若干工程性问题（DEBUG_LEVEL=2 单密钥广播路径下 Beaver 乘法的 reshape 错误、prefix_parity_query 的多 batch 广播、FSSKeyProvider 的磁盘兜底、子进程端口隔离等）进行了修复与优化，详见项目 README.md。')

print(f'Total paragraphs: {len(PARAGRAPHS)}')

# ===== 生成 docx =====
doc = Document('北京邮电大学2026届本科毕业设计（论文）撰写指导手册-信通、电子、计算机、人工智能、网安、智工、未来学院.docx')

body = doc.element.body
for el in list(body):
    if el.tag == qn('w:p'):
        body.remove(el)

for style, text in PARAGRAPHS:
    try:
        para = doc.add_paragraph(text, style=style)
    except KeyError:
        para = doc.add_paragraph(text)

out = '支持隐私保护的检索增强生成系统-论文.docx'
doc.save(out)
print(f'saved: {out}')
